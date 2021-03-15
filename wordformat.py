# module for docxtest
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

# Links to an external hyperlink
def add_hyperlink(paragraph, text, url):
    # Gets access to the document.xml.rels file and add new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement("w:bookmarkStart")
    start.set(docx.oxml.ns.qn("w:id"), "0")
    start.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(start)
    text = docx.oxml.OxmlElement("w:r")
    text.text = bookmark_text
    tag.append(text)
    end = docx.oxml.shared.OxmlElement("w:bookmarkEnd")
    end.set(docx.oxml.ns.qn("w:id"), "0")
    end.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(end)
    return run


doc = docx.Document("docs/Cyber_Clips_February_23_2021.docx")

# Adds text to paragraph with link to an internal bookmark
# Returns run with link in it
def book_link(paragraph, bookmark_name, text, tool_tip=None):
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("w:anchor"), bookmark_name)
    if tool_tip:
        hyperlink.set(docx.oxml.shared.qn("w:tooltip"), tool_tip)
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.name = "Times New Roman"
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return r
