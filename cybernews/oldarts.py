import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
import re
from datetime import datetime
from dateutil import parser
import json
import csv
import logging
import traceback
from spiders import articlesspider as AS
from scrapy.crawler import CrawlerProcess
from scrapy.crawler import CrawlerRunner
from scrapy.utils.project import get_project_settings
from itemadapter import ItemAdapter
from itertools import groupby
import win32com.client as win32

Febpath = "C:/Users/tlebr/OneDrive - DEFDEM/Theo Lebryk/Daily Clippings/February/"
Marpath = "C:/Users/tlebr/OneDrive - DEFDEM/Theo Lebryk/Daily Clippings/March/"
path2016 = r"C:\Users\tlebr\OneDrive - DEFDEM\Theo Lebryk\OneDrive_2021-04-01\Research - Daily Clips\2017/"

logging.basicConfig(
    filename="oldarts.log", encoding="utf-8", filemode="w", level=logging.DEBUG
)
# returns list of tuples (url, ranking, filename)
def get_all_links(path):
    links = []
    for filename in os.listdir(path):
        f = path + filename
        with zipfile.ZipFile(f) as z:
            d = z.read("word/document.xml")
        xml_str = str(d)
        # [1:] gets rid of microsoft random urls
        link_list = re.findall(r"http.*?\<", xml_str)[1:]
        # strips trailing ">" from urls
        link_list = [(x[:-1], i + 1, filename) for i, x in enumerate(link_list)]
        links += link_list
    return links

def get_doc_links(docpath):
    with zipfile.ZipFile(docpath) as z:
        d = z.read("word/document.xml")
    xml_str = str(d)
    # [1:] gets rid of microsoft random urls
    link_list = re.findall(r"<w:t>http.*?\<", xml_str)
    # strips leading <w:t> and trailing ">" from urls
    link_list = [x[5:-1] for x in link_list]
    # links += link_list
    return link_list

def get_headline(d):
    headlines = []
    for paragraph in d.paragraphs:
        if (paragraph.style.name == "Heading 1" and 
                paragraph.text not in headlines and
                len(paragraph.text) > 20):
            headlines.append(paragraph.text)
    return headlines


def get_meta(d):
    result = []
    for i, para in enumerate(d.paragraphs):
        # line after a heading that is not another heading and not blank line 
        # should contain meta data (author1, author2 ..., source, date, )
        if (para.style.name == "Heading 1" and i+1 < len(d.paragraphs)):
            if (d.paragraphs[i+1].style.name != "Heading 1" and
                    len(d.paragraphs[i+1].text) > 10 and
                    len(d.paragraphs[i + 1].text.split(",")) > 3):
                meta = d.paragraphs[i + 1].text
                m_list = re.split(", | and ", meta)
                dt = ",".join(m_list[-2:]).strip()
                try: 
                    dt = parser.parse(dt)
                except Exception as e:
                    logging.error(f"{filename} date error")
                    logging.debug(e)
                    logging.debug(f'dt is {dt}')
                dct = {
                    # last two elements in "month day, year" format
                    # to export to json, convert to date then reconvert to string
                    'date' : str(dt),
                    # thrid to last element should be source
                    'source' : m_list[-3],
                    'author' : m_list[:-3],
                }
                result.append(dct)
    return result

def get_body(d):
    result = []
    for i, para in enumerate(d.paragraphs):
        # j will use i+2 to check for body after meta. 
        if para.style.name == "Heading 1" and i+2 < len(d.paragraphs):
            if (d.paragraphs[i+1].style.name != "Heading 1"
                        and len(d.paragraphs[i+1].text) > 15):
                # count paragraphs before next header
                j = i + 2
                while d.paragraphs[j].style.name != "Heading 1":
                    j += 1
                    if j >= len(d.paragraphs):
                        break
                # body should have more than one paragraph before next header
                # and should not start with "("
                if j - i >= 4:
                    if d.paragraphs[i+2].text.strip():
                        if d.paragraphs[i+2].text[0] == "(":
                            continue
                    j = i + 2
                    body = ""
                    while d.paragraphs[j].style.name != "Heading 1":
                        body += d.paragraphs[j].text
                        j+=1
                        if j >= len(d.paragraphs):
                            break
                    if len(body) < 65:
                        continue
                    result.append(body)
    return result

if __name__ == "__main__":
    articles = []
    path = path2016
    error_count = 0
    for filename in os.listdir(path):
        f = path + filename
        links = get_doc_links(f)
        d = Document(f)
        body = get_body(d)
        meta = get_meta(d)
        headline = get_headline(d)
        # populate missing headlines with empty string
        if len(links) > len(headline):
            headline += ["" for i in range(len(links) - len(headline))]
        if len(links) < len(body):
            s = f"links length: {len(links)}, \n\
body length: {len(body)}, \n\
meta length: {len(meta)}"
            logging.error(f"{filename} lens mismatch")
            logging.debug(s)
            logging.debug(f'meta: {meta} \nlinks: {links}')
            logging.debug(f'body: {body[:100]}')
            error_count += 1
            continue
        for i in range(len(links)):
            el = {
            "title": headline[i],
            "url": links[i],
            "ranking": i+1,
            "filename": filename,
            "body": body[i],
            }
            if len(body) == len(meta):
                el.update(meta[i])
            articles.append(el)
    # with open(f'C:/Users/tlebr/Google Drive/fdd/dailynews/cybernews/data/jsons/{path.split(r"/")[-2]}.json', 'w', encoding="utf8") as fout:

    with open(f'C:/Users/tlebr/Google Drive/fdd/dailynews/cybernews/data/jsons/2017.json', 'w', encoding="utf8") as fout:
        json.dump(articles, fout, ensure_ascii=False, indent=1)

    logging.info(f"Error rate: {error_count/len(os.listdir(path))}%")

    # settings = get_project_settings()
    # settings["FEEDS"] = {
    #     "test1.json": {"format": "json", "encoding": "utf8", "overwrite": False}
    # }
    # settings["LOG_LEVEL"] = "INFO"
    # links = get_all_links(path=Febpath) + get_all_links(path=Marpath)
    # process = AS.get_articles(links, settings)
    # process.start()
    # datapath = "C:\\Users\\tlebr\\Google Drive\\fdd\\dailynews\\data\\"
    # with open(datapath + "oldarts.json", "w", encoding="utf-8") as fp:
    #     s = json.dumps(AS.DICT_LS, ensure_ascii=False, indent="\t")
    #     fp.write(s)
    # with open(datapath + "oldarts.csv", "w", encoding="utf-8-sig") as fp:
    #     writer = csv.DictWriter(fp, fieldnames=AS.DICT_LS[0].keys())
    #     writer.writeheader()
    #     for data in AS.DICT_LS:
    #         writer.writerow(data)
