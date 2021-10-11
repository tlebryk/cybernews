"""Script to export daily briefing into a word document"""

import json
import sys
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor, Length, Pt
from docx import Document
from wordformat import add_hyperlink, add_bookmark, book_link

# briefing should be either "CEFP News Briefing"
# or "Daily Cyber News Briefing"
def exporter(briefing, export_path, obj):
    doc = Document()
    styles = doc.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    font = new_heading_style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0,0,0)

    normal = styles["Normal"]
    font = normal.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    para = doc.add_paragraph()
    para.add_run(briefing).bold = True
    para = doc.add_paragraph()
    para.add_run(datetime.today().strftime("%A, %B %d, %Y")).bold = True
    para = doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("Contents").bold = True
    # save space for internal document links/bookmarks
    parals = [doc.add_paragraph() for i in range(len(obj))]
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("Articles").bold = True

    # flow in articles with bookmarks in headings
    for i, o in enumerate(obj):
        h1 = doc.add_heading()
        run1 = add_bookmark(h1, o.title, str(i))
        run1.bold = True
        font = run1.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0,0,0)
        para2 = doc.add_paragraph()
        para2.add_run(o.authors + ", ")
        para2.add_run(o.source).italic = True
        dt = o.briefingdate
        para2.add_run(", " + str(dt))
        para3 = doc.add_paragraph()
        add_hyperlink(para3, o.url, o.url)
        doc.add_paragraph()
        para4 = doc.add_paragraph()
        para4.add_run(o.body)
        doc.add_paragraph()
        doc.add_paragraph()

    for i, o in enumerate(obj):
        r = book_link(paragraph=parals[i], bookmark_name=str(i), text=o.title, tool_tip=o.title)
        r.bold = True
        parals[i].add_run(' (')
        parals[i].add_run(o.source).italic = True
        parals[i].add_run(')')


    for para in doc.paragraphs:
        para_format = para.paragraph_format
        para_format.line_spacing = 1
        para_format.space_before = 0 
        para_format.space_after = 0
    doc.save(export_path)


def CEFP_export(export_path, obj):
    exporter("CEFP News Briefing", export_path,obj) 

def cyber_export(export_path, obj):
    exporter("Daily Cyber News Briefing", export_path, obj)


if __name__ == '__main__':
    path = sys.argv[1]
    with open(path, 'r', encoding='utf-8') as f:
        data=f.read()
    export_path = sys.argv[2]
    cyber_export(export_path, data)


